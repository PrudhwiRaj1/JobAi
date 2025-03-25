from django.conf import settings
from django.shortcuts import get_object_or_404
from django.http.response import JsonResponse
from django.contrib import messages
from django.core.mail.message import EmailMessage
import os
import json
import re
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
import comtypes  # Windows only (Handled below)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, Image, SimpleDocTemplate
from reportlab.lib import colors  # Fixed import
from .models import JobApplication, jobseeker_profile

def convert_docx_to_json(docx_path, filename):
    """Convert a .docx file to JSON format and save it to media folder"""
    document = Document(docx_path)
    data = {"paragraphs": []}

    for para in document.paragraphs:
        data["paragraphs"].append(para.text)
    
    json_path = os.path.join(settings.MEDIA_ROOT, "json", filename + ".json")
    
    try:
        with open(json_path, "w") as json_file:
            json.dump(data, json_file, indent=4)
    except Exception as e:
        print(f"Error saving JSON file: {e}")
    
    return json.dumps(data, indent=4)
# Create your views here.

def convert_doc_to_docx(doc_path):
    """Converts a DOC file to DOCX using comtypes (Windows only)."""
    word = comtypes.client.CreateObject('Word.Application')#pylint:disable=undefined-variable
    doc = word.Documents.Open(doc_path)
    docx_path = doc_path + "x"
    doc.SaveAs(docx_path, FileFormat=16)  # 16 = docx format
    doc.Close()
    word.Quit()
    return docx_path

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_resume_text(profile_id):
    """Fetches the resume path from the database and extracts text."""
    jobseeker = get_object_or_404(jobseeker_profile, id=profile_id)
    resume_path = os.path.join(settings.MEDIA_ROOT, str(jobseeker.resume))
    
    if resume_path.endswith(".docx"):
        return extract_text_from_docx(resume_path)
    elif resume_path.endswith(".doc"):
        docx_path = convert_doc_to_docx(resume_path)
        return extract_text_from_docx(docx_path)
    
    return None

def calculate_ats_score(resume_text, job_description):
    prompt = f"""
    You are an ATS resume screening AI. Evaluate the resume against the job description.

    **Scoring Criteria**:
    - **Skills and Qualifications match(50%)**
    - **Formatting & Readability (15%)**
    - **Grammar & Clarity (40%)**
    
    minimum 1/3 of appliication must be accepted so increase ats score of 3 for same job
    
    **Resume**:
    {resume_text}

    **Job Description**:
    {job_description}
    
    Provide:
    1. **ATS Score** (out of 100)
    2. **Improvement Suggestions** (if score < 50)
    """

    response = openai.ChatCompletion.create( #pylint: disable=undefined-variable
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    
    ats_result = response["choices"][0]["message"]["content"]
    ats_score = int(re.search(r'\b(\d{1,3})\b', ats_result).group(1))
    return ats_score

def generate_admit_card(jobseeker, job_title, application_id):
    """Generates an Admit Card PDF with jobseeker details, company logo, and profile image."""
    
    # File Path for PDF
    file_name = f"{jobseeker.name.replace(' ', '_')}_admit_card.pdf"
    file_path = os.path.join(settings.MEDIA_ROOT, "admit_cards", file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    application = get_object_or_404(JobApplication, id=application_id)

    # Create PDF Document
    pdf = SimpleDocTemplate(file_path, pagesize=letter)  # ✅ Now this works
    elements = []
    styles = getSampleStyleSheet()

    # Title
    title = Paragraph("<b>Aptitude Test Admit Card</b>", styles['Title'])
    
    # Load Company Logo
    logo_path = application.company_joblist.company.profile_img.path if application.company_joblist.company.profile_img else None
    company_logo = Image(logo_path, width=120, height=60) if logo_path and os.path.exists(logo_path) else Paragraph("<b>[Company Logo Not Available]</b>", styles['Normal'])
    
    # Load Jobseeker Profile Image
    profile_path = jobseeker.profile_img.path if jobseeker.profile_img else None
    applicant_image = Image(profile_path, width=100, height=100) if profile_path and os.path.exists(profile_path) else Paragraph("<b>[Applicant Image Not Available]</b>", styles['Normal'])
    
    # Table for Header with Company Logo & Title Alignment
    header_table = Table([[company_logo, title]], colWidths=[130, 400])
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),  # Align logo left
        ('ALIGN', (1, 0), (1, 0), 'CENTER')  # Align title center
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 20))  # Space between title and next section
    
    # Jobseeker Image Alignment (Top Right)
    applicant_image_table = Table([[applicant_image]], colWidths=[100], rowHeights=[100])
    applicant_image_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Align jobseeker image to the right
    ]))
    elements.append(applicant_image_table)
    elements.append(Spacer(1, 20))
    
    # Jobseeker Information Table
    data = [
        ["Applicant Name:", jobseeker.name],
        ["Email:", jobseeker.email],
        ["Phone:", jobseeker.phone],
        ["Qualification:", jobseeker.highest_qualification],
        ["Job Title:", job_title],
        ["Company:", application.company_joblist.company.name],
        ["Status:", "Selected"]
    ]

    table = Table(data, colWidths=[150, 300])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    elements.append(table)
    
    # Final Instructions
    instructions = Paragraph("""
        <br/><b>Instructions:</b><br/>
        Please bring this admit card along with required documents on the test day.
    """, styles['Normal'])
    elements.append(instructions)
    
    # Generate PDF
    pdf.build(elements)
    
    return file_path

def auto_process_application(request, application_id):
    """Automatically processes the job application based on ATS score."""
    try:
        application = get_object_or_404(JobApplication, id=application_id)
        jobseeker = get_object_or_404(jobseeker_profile, id=application.jobseeker_id)
        resume_text = extract_resume_text(application.jobseeker_id)
        if not resume_text:
            return {"error": "Failed to extract resume text."}
        else:
            print(resume_text)
        # Skip processing if already completed
        if application.ats_score is not None:
            return JsonResponse({
                "status": application.status,
                "ats_score": application.ats_score
            })

        #  Calculate ATS Score
        ats_score = calculate_ats_score(resume_text,application.company_joblist.job_description)


        #  Save ATS Score and Status in Database
        application.ats_score = ats_score
        application.save()

        # Determine Acceptance or Rejection
        if ats_score >= 55:
            application.status = "Accepted"
            status_text = "Accepted"
            status_class = "text-success fw-bold"
            
            # Generate Admit Card
            admit_card_path = generate_admit_card(jobseeker, application.company_joblist.job_title.job_title,application_id)
            
            # Email Content for Accepted Applicants
            email_subject = f"You're eligible to attend for Aptitude test for position {application.company_joblist.job_title.job_title} at {application.company_joblist.company.name}"
            email_body = f"""
Dear {jobseeker.name},

We are thrilled to inform you that your application for the position of {application.company_joblist.job_title.job_title} has been Accepted! 
 
Your admit card is attached with this email.

Next Steps:  
- Kindly review the admit card and take print out.
- Keep your admit card with you while come for aptitude test. 
- Ensure your participation in aptitude test.
 

We hope you prepare well for the test .The more details regarding Date,Time,Venue of exam will be informed shortly.

Best Regards,  
Recruitment Team
"""
            email = EmailMessage(email_subject, email_body, "noreply@company.com", [jobseeker.email])
            email.attach_file(admit_card_path)
            email.send()

        else:
            application.status = "Rejected"
            status_text = "Rejected"
            status_class = "text-danger fw-bold"

            # Email Content for Rejected Applicants
            email_subject = f"⚠️ Application Update - {application.company_joblist.job_title.job_title}"
            email_body = f"""
Dear {jobseeker.name},

We regret to inform you that your application for the position of {application.company_joblist.job_title.job_title} has been Rejected.  

We encourage you to improve your resume and apply for future opportunities with us.

Also wait for further mails as some times(depends on some criteria and company decision ) even the rejected list may also considered to next level.

Wishing you all the best in your job search!

Best Regards,  
Recruitment Team
"""
            email = EmailMessage(email_subject, email_body, "noreply@company.com", [jobseeker.email])
            email.send()

        # Save ATS Score & Status
        application.ats_score = ats_score
        application.save()
        if ats_score>100:
            messages.error(request,"Error occurred")
        # Return JSON Response for Frontend Update
        return JsonResponse({
            "status": application.status,
            "status_text": status_text,
            "status_class": status_class,
            "ats_score": ats_score
        })

    except JobApplication.DoesNotExist:
        return JsonResponse({"error": "Application not found"}, status=404)